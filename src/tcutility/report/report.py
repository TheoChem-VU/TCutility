import os
import pathlib as pl
import uuid
from importlib.util import find_spec
from typing import TYPE_CHECKING, List, Tuple, Union

from tcutility import errors

if find_spec("docx") is None:
    raise errors.MissingOptionalPackageError("docx")
if find_spec("htmldocx") is None:
    raise errors.MissingOptionalPackageError("htmldocx")
if find_spec("cv2") is None:
    raise errors.MissingOptionalPackageError("opencv-python")


import cv2  # noqa: E402 # This is the opencv-python package
import docx
import htmldocx
import numpy as np
import PIL
import PIL.Image
import PIL.ImageDraw
import PIL.ImageFont
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt
from htmldocx import HtmlToDocx

from tcutility.report.formatters.generic import WordFormatter
from tcutility.report.formatters.xyz import StandardXYZFormatter
from tcutility.results.read import read
from tcutility.results.result import Result

if TYPE_CHECKING:
    import docx.document
    import docx.oxml.ns
    import docx.oxml.shared


def _add_page_numbers(document: "docx.document.Document") -> None:
    document.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = document.sections[0].footer.paragraphs[0].add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"  # type: ignore  # text not found in docx by mypy

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _set_cell_border(cell, **kwargs):
    """
    Set a cell`s border
    Usage:

    _set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ("start", "top", "end", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))  # type: ignore
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn("w:{}".format(key)), str(edge_data[key]))


def _set_repeat_table_header(row):
    """set repeat table row on every new page"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    return row


def _set_cell_color(cell, color):
    shading_elm_1 = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


table_formatting = Result()
table_formatting.table.alignment = WD_TABLE_ALIGNMENT.CENTER  # type: ignore  # WD_TABLE_ALIGNMENT not found in docx by mypy
table_formatting.font.size = Pt(10.5)  # type: ignore  # Pt not found in docx by mypy


class DocxFigure:
    def __init__(self, doc: docx.document.Document, width=None, height=None):
        self.caption = ""
        self.doc = doc
        self.width = width or Cm(17.8)
        self.height = height

        self.figures = []
        self.html_parser = htmldocx.HtmlToDocx()

    def __enter__(self):
        return self

    def __exit__(self, *args):
        self.write()

    def add_image(self, path: str, row: Union[int, Tuple[int]], col: Union[int, Tuple[int], None] = None, label: str = ""):
        img = cv2.imread(path)

        self.figures.append(
            {
                "path": path,
                "img": img,
                # 'col': (col, col) if isinstance(col, int) else col,
                # 'row': (row, row) if isinstance(row, int) else row,
                "col": col,
                "row": row,
                "label": label,
            }
        )

    def write(self):
        new_img = self._stitch_images()
        new_img_path = str(uuid.uuid4()) + ".png"
        cv2.imwrite(new_img_path, new_img)

        p = self.doc.add_paragraph()
        r = p.add_run()
        r.add_picture(new_img_path, width=self.width, height=self.height)

        self._write_caption()

        os.remove(new_img_path)

    def _stitch_images(self):
        grid_width, grid_height = max(fig["col"] for fig in self.figures) + 1, max(fig["row"] for fig in self.figures) + 1
        col_widths = [max(fig["img"].shape[1] for fig in self.figures if fig["col"] == index) for index in range(grid_width)]
        dpi = sum(col_widths) / (self.width / 360000 * 0.393700787)
        letter_extra_rows = int(dpi * 12 / 72) + 1

        row_heights = [max(fig["img"].shape[0] + letter_extra_rows for fig in self.figures if fig["row"] == index) for index in range(grid_height)]

        for fig in self.figures:
            height = row_heights[fig["row"]]
            width = col_widths[fig["col"]]

            fig["img"] = cv2.copyMakeBorder(fig["img"], height - fig["img"].shape[0], 0, 0, width - fig["img"].shape[1], cv2.BORDER_CONSTANT, value=[255, 255, 255])

            font_path = "Times New Roman"
            font_size = letter_extra_rows

            font = PIL.ImageFont.truetype(font_path, font_size)
            image_pil = PIL.Image.fromarray(fig["img"])  # Convert OpenCV image to PIL image
            draw = PIL.ImageDraw.Draw(image_pil)
            if fig["label"]:
                draw.text((0, 0), fig["label"], font=font, fill=(0, 0, 0))
            fig["img"] = np.array(image_pil)  # Convert PIL image back to OpenCV image

        new_img = np.zeros((sum(row_heights), sum(col_widths), 3))

        print(self.width / 360000 * 0.393700787, Inches(self.width), sum(col_widths), dpi)

        for fig in self.figures:
            start_pos = sum(row_heights[: fig["row"]]), sum(col_widths[: fig["col"]])
            end_pos = sum(row_heights[: fig["row"] + 1]), sum(col_widths[: fig["col"] + 1])
            new_img[start_pos[0] : end_pos[0], start_pos[1] : end_pos[1], :] = fig["img"]

        return new_img

    def _write_caption(self):
        paragraph = self.doc.add_paragraph("Figure S")
        paragraph.runs[-1].bold = True

        # numbering field
        run = paragraph.add_run("")
        run.bold = True

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.text = " SEQ Figure \\* ARABIC"  # type: ignore  # text not found in docx by mypy
        run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)

        html_doc = self.html_parser.parse_html_string(f"<b>.</b> {self.caption}")

        for add_run in html_doc.paragraphs[-1].runs:
            run = paragraph.add_run(add_run.text)
            run.bold = add_run.bold
            run.italic = add_run.italic
            run.underline = add_run.underline  # type: ignore  # underline not found in docx by mypy
            run.font.superscript = add_run.font.superscript
            run.font.subscript = add_run.font.subscript
            run.font.strike = add_run.font.strike
            run.font.size = add_run.font.size


class DocxTable:
    def __init__(self, file: Union[str, pl.Path, docx.document.Document] = "test.docx"):
        self.file = file
        self.dont_save = False
        if isinstance(file, docx.document.Document):
            self.dont_save = True
            self.doc = file
        else:
            if not os.path.exists(file):
                self.doc = docx.Document()
            else:
                self.doc = docx.Document(str(file))

        self.caption = ""
        self.columns = []
        self.column_options = []
        self.footnote = None
        self.rows = []
        self.mergers = []
        self.html_parser = htmldocx.HtmlToDocx()

    def __enter__(self):
        return self

    def __exit__(self, *args):
        self.write()

    def add_column(self, name, **kwargs):
        self.columns.append(["single", name, kwargs])
        self.column_options.append(kwargs)

    def add_column_group(self, group_name, column_names, **kwargs):
        self.columns.append(["grouped", group_name, column_names])
        [self.column_options.append(kwargs) for _ in column_names]

    def add_row(self, data):
        self.rows.append(["data", data])

    def add_header_row(self, name):
        self.rows.append(["header", name])

    def add_empty_row(self):
        self.rows.append(["empty"])

    def merge_cells(self, x, y):
        if isinstance(x, int):
            x = (x, x)
        if isinstance(y, int):
            y = (y, y)

        self.mergers.append([x, y])

    def _correct_table_dims(self):
        num_cols = 2 * len([col for col in self.columns if col[0] == "single"]) + sum([len(col[2]) + 1 for col in self.columns if col[0] == "grouped"]) - 1
        num_rows = 2 + len(self.rows)

        for _ in range(num_cols - len(self.tab.columns)):
            self.tab.add_column(Cm(0.1))
        for _ in range(num_rows - len(self.tab.rows)):
            self.tab.add_row()

    @property
    def shape(self):
        Ncols = 0

        for col in self.columns:
            if col[0] == "single":
                Ncols += 2
            if col[0] == "grouped":
                Ncols += len(col[2]) + 1

        return 2 + len(self.rows), Ncols - 1

    def write(self):
        self._write_caption()
        # write the table caption
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.tab = self.doc.add_table(1, 1)
        self._correct_table_dims()

        for x, y in self.mergers:
            self.tab.cell(x[0], y[0]).merge(self.tab.cell(x[1], y[1]))

        num_cols = 2 * len([col for col in self.columns if col[0] == "single"]) + sum([len(col[2]) + 1 for col in self.columns if col[0] == "grouped"]) - 1
        # create a table
        self.tab.alignment = table_formatting.table.alignment  # type: ignore # Results object does not have typing

        # write the column headers
        spacing_columns = []
        col_idx = 0
        for col in self.columns:
            if col[0] == "single":
                self.write_cell(1, col_idx, col[1], bold=True, font_size=table_formatting.font.size)  # type: ignore  # results object does not have typing
                spacing_columns.append(col_idx + 1)
                col_idx += 2

            if col[0] == "grouped":
                _ = self.write_cell(0, (col_idx, col_idx + len(col[2]) - 1), col[1], bold=True, bottom={"sz": 12, "val": "single", "color": "#000000"}, font_size=table_formatting.font.size)  # type: ignore  # results object does not have typing

                for i, val in enumerate(col[2]):
                    self.write_cell(1, col_idx + i, val, bold=True, font_size=table_formatting.font.size)  # type: ignore  # results object does not have typing

                spacing_columns.append(col_idx + len(col[2]))
                col_idx += len(col[2]) + 1

        # set the lines for the top and bottom header rows
        for i in range(num_cols):
            _set_cell_border(self.tab.cell(0, i), top={"sz": 12, "val": "single", "color": "#000000"})
            _set_cell_border(self.tab.cell(1, i), bottom={"sz": 12, "val": "single", "color": "#000000"})

        for j, row in enumerate(self.rows):
            if row[0] == "data":
                for i in range(num_cols):
                    if i in spacing_columns:
                        continue
                    num_spacing_past = len([k for k in spacing_columns if (k - 1) < i])
                    self.write_cell(j + 2, i, row[1][i - num_spacing_past], font_size=table_formatting.font.size, **self.column_options[i - num_spacing_past])  # type: ignore  # results object does not have typing

            if row[0] == "header":
                _ = self.write_cell(
                    j + 2,
                    (0, num_cols - 1),
                    row[1],
                    bold=True,
                    top={"sz": 12, "val": "single", "color": "#000000"},
                    bottom={"sz": 12, "val": "single", "color": "#000000"},
                    bkgr_color="F2F2F2",
                    font_size=table_formatting.font.size,  # type: ignore  # font not found in docx by mypy
                )

            if row[0] == "empty":
                for i in range(num_cols):
                    self.write_cell(j + 2, i, "")

        _set_repeat_table_header(self.tab.rows[0])
        _set_repeat_table_header(self.tab.rows[1])

        self._write_footnote()

        if not self.dont_save:
            self.doc.save(str(self.file))

    def write_cell(self, row, col, text, alignment="center", vert_alignment="center", bold=None, italic=None, bkgr_color=None, font_size=None, **kwargs):
        if isinstance(row, int) and isinstance(col, int):
            cell = self.tab.cell(row, col)
        else:
            if isinstance(row, int):
                row = (row, row)
            if isinstance(col, int):
                col = (col, col)
            cell = self.tab.cell(row[0], col[0]).merge(self.tab.cell(row[1], col[1]))

        if len(cell.paragraphs[0].runs) > 0:
            return

        # cell.text = str(text).strip()
        alignment = {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
        }[alignment]

        vert_alignment = {
            "center": WD_ALIGN_VERTICAL.CENTER,
            "top": WD_ALIGN_VERTICAL.TOP,
            "bottom": WD_ALIGN_VERTICAL.BOTTOM,
        }[vert_alignment]

        # for txt, settings in parse_text(text):
        self.html_parser.add_html_to_cell(text.replace("-", "–"), cell)

        cell.paragraphs[0].alignment = alignment
        cell.vertical_alignment = vert_alignment
        for run in cell.paragraphs[0].runs:
            if bold:
                run.bold = bold
            if italic:
                run.italic = italic
            if font_size:
                run.font.size = font_size

        _set_cell_border(cell, **kwargs)
        if bkgr_color is not None:
            _set_cell_color(cell, bkgr_color)

        return cell

    def add_footnote(self, text, bold=None, italic=None, font_size=None):
        self.footnote = (text, bold, italic, font_size)

    def _write_footnote(self):
        if self.footnote is None:
            return

        # set the lower line
        for col in range(self.shape[1]):
            cell = self.tab.cell(self.shape[0] - 1, col)
            _set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "#000000"})

        self.html_parser.add_html_to_document(self.footnote[0], self.doc)

        for run in self.doc.paragraphs[0].runs:
            if self.footnote[1]:
                run.bold = self.footnote[1]
            if self.footnote[2]:
                run.italic = self.footnote[2]
            if self.footnote[3]:
                run.font.size = self.footnote[3]

    def _write_caption(self):
        paragraph = self.doc.add_paragraph("Table S")
        paragraph.runs[-1].bold = True

        # numbering field
        run = paragraph.add_run("")
        run.bold = True

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.text = " SEQ Table \\* ARABIC"  # type: ignore  # text not found in docx by mypy
        run._r.append(instrText)

        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)

        html_doc = self.html_parser.parse_html_string(f"<b>.</b> {self.caption}")

        for add_run in html_doc.paragraphs[-1].runs:
            run = paragraph.add_run(add_run.text)
            run.bold = add_run.bold
            run.italic = add_run.italic
            run.underline = add_run.underline  # type: ignore  # underline not found in docx by mypy
            run.font.superscript = add_run.font.superscript
            run.font.subscript = add_run.font.subscript
            run.font.strike = add_run.font.strike
            run.font.size = add_run.font.size


class SI:
    def __init__(self, path: Union[str, pl.Path], overwrite: bool = False, font: str = "Arial") -> None:
        """Initializes the SI class for creating supporting information (SI) files in Microsoft Word format.

        This class is responsible for creating and managing a Microsoft Word document that serves as supporting information (SI) for reports or publications.
        It allows for the addition of various elements such as text, headings, and formatted content from HTML.

        Args:
            path (str | pl.Path): The location of the Word file. Does not have to have a file-extension.
            overwrite (bool, optional): Whether to append to or overwrite the file. Defaults to False.
            font (str, optional): The font to be used in the document. Defaults to "Arial".

        Attributes:
            path (pl.Path): The path to the Word document.
            doc (docx.Document): The Word document object.
        """
        self.path = pl.Path(path).with_suffix(".docx")
        self.doc = docx.Document()

        if not os.path.exists(self.path) or overwrite:
            self.doc = docx.Document()
        else:
            self.doc = docx.Document(str(self.path))

        self.doc.styles["Normal"].font.name = "Times New Roman"  # type: ignore  # font not found in docx by mypy
        self.doc.styles["Normal"].font.size = Pt(12)  # type: ignore  # font not found in docx by mypy
        self.doc.styles["Normal"].paragraph_format.space_after = 1  # type: ignore  # paragraph_format not found in docx by mypy
        self.html_parser = htmldocx.HtmlToDocx()  # type: ignore  # HtmlToDocx not found in docx by mypy

        # Set the font to the specified font
        self.doc.styles["Normal"].font.name = font  # type: ignore  # font not found in docx by mypy

    def __enter__(self):
        """Enables the use of the class as a context manager."""
        return self

    def __exit__(self, *args):
        for section in self.doc.sections:
            # 1.9 cm
            section.left_margin = Cm(1.9)
            section.right_margin = Cm(1.9)

        _add_page_numbers(self.doc)

        self.doc.save(str(self.path))

    def add_figure(self, *args, **kwargs) -> DocxFigure:
        return DocxFigure(self.doc, *args, **kwargs)

    def add_table(self) -> DocxTable:
        return DocxTable(self.doc)

    def add_xyz(self, obj: Union[str, Result, pl.Path], title: Union[str, None] = None, formatter: WordFormatter = StandardXYZFormatter()) -> None:
        """Adds XYZ formatted content to the document.

        This method is responsible for adding the coordinates and information about a calculation to the supporting information document.
        It includes details such as the electronic bond energy, Gibb's free energy, enthalpy, imaginary mode, and the coordinates of the molecule.

        Args:
            obj: A string specifying a calculation directory or a `TCutility.results.Result` object from a calculation.
            title: The title to be written before the coordinates and information. If None, no title is added.
            formatter: The formatter to be used for formatting the content. Defaults to `StandardXYZFormatter`.

        Returns:
            None
        """
        ret_str = ""

        # Add the formatted content to the document
        if isinstance(obj, str) or isinstance(obj, pl.Path):
            calc_path = pl.Path(obj)
            if not calc_path.is_dir():
                return  # raise ValueError(f"Invalid calculation directory: {calc_path}")
            obj = read(obj)

        ret_str += formatter.format(obj, title=title)

        # print(ret_str)
        parser = HtmlToDocx()
        parser.add_html_to_document(ret_str, self.doc)
        return

    def add_heading(self, text: str, level: int = 1) -> None:
        """Adds a heading to the document.

        This method allows for the addition of a heading to the Word document, with customizable text and level.

        Args:
            text (str): The text of the heading.
            level (int, optional): The level of the heading (determines the size and style). Defaults to 1.
        """
        self.doc.add_heading(text, level)

    def add_page_break(self) -> None:
        self.doc.add_page_break()

    def add_toc(self) -> None:
        title = self.doc.add_paragraph()
        title_run = title.add_run("Contents")
        title_run.bold = True
        title_run.font.size = Pt(14)
        title.paragraph_format.space_after = Pt(10)

        # figures
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        fldChar.set(qn("w:dirty"), "true")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'TOC \\h \\z \\c "Figure"'  # type: ignore # "Table" of list of table and "Figure" for list of figure
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:t")
        fldChar3.text = "Right-click to update field."  # type: ignore # text not found in docx by mypy
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar4)

        # tables
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        fldChar.set(qn("w:dirty"), "true")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'TOC \\h \\z \\c "Table"'  # type: ignore # "Table" of list of table and "Figure" for list of figure
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        fldChar3 = OxmlElement("w:t")
        fldChar3.text = "Right-click to update field."  # type: ignore # text not found in docx by mypy
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar4)


def get_subdirs(root_folder: pl.Path) -> List[pl.Path]:
    """Iteratively searches through a folder and returns all the most nested subdirs."""
    most_nested_subdirs = []
    for root, dirs, files in os.walk(root_folder):
        # If 'dirs' is empty, it means 'root' contains no subdirectories, thus it is most nested.
        if not dirs:
            most_nested_subdirs.append(pl.Path(root))
    return most_nested_subdirs


def replace_files_rkf_to_ams_rkf(root_folder: pl.Path) -> None:
    """Iteratively searches through a folder and replaces all files with the extension '.rkf' to '.ams.rkf', except if the file has 'adf.rkf' in the name."""
    for root, dirs, files in os.walk(root_folder):
        for file in files:
            if file.endswith(".rkf") and "adf.rkf" not in file:
                new_name = file.replace(".rkf", ".ams.rkf")
                os.rename(pl.Path(root) / file, pl.Path(root) / new_name)


def main():
    calc_dir = pl.Path("__file__").resolve().parents[0] / "test" / "fixtures"
    main_path = pl.Path("__file__").resolve().parents[0] / "examples"

    # all_subdirs = [folder for folder in calc_dir.iterdir() if folder.is_dir()]
    all_subdirs = get_subdirs(calc_dir)
    res_objects = []
    for dir_ in all_subdirs:
        try:  # Try to read the results of the calculation
            res_objects.append(read(dir_))
        except Exception:
            pass

    with SI(main_path / "test", overwrite=False) as si:
        si.add_heading("SI project")
        for obj in res_objects:
            si.add_xyz(obj=obj)


if __name__ == "__main__":
    main()
