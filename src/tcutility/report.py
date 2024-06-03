import pathlib as pl
from enum import Enum

import docx
from htmldocx import HtmlToDocx

from tcutility import results
from tcutility.results import Result

from tcutility._report.formatters import xyz

def _hydrogen_bond_format(obj: Result, write_str: str) -> str:
    # add electronic energy. E should be bold and italics. Unit will be kcal mol^-1
    E = str(round(obj.properties.energy.bond, 2)).replace("-", "–")
    write_str += f"<b><i>E</i></b> = {E} kcal mol<sup>—1</sup><br>"

    # add Gibbs and enthalpy if we have them
    if obj.properties.energy.enthalpy:
        H = str(round(obj.properties.energy.enthalpy, 2)).replace("-", "–")
        write_str += f"<b><i>H</i></b> = {H} kcal mol<sup>—1</sup><br>"
    if obj.properties.energy.gibbs:
        G = str(round(obj.properties.energy.gibbs, 2)).replace("-", "–")
        write_str += f"<b><i>G</i></b> = {G} kcal mol<sup>—1</sup><br>"

    # add imaginary frequency if we have one
    if obj.properties.vibrations:
        num_imag_modes = obj.properties.vibrations.number_of_imag_modes

        # We only want to write nimag = 0 if there are no imaginary modes
        if num_imag_modes < 1:
            return write_str + "<b><i>ν<sub>imag</sub></i></b> = 0"

        # Otherwise we write the imaginary frequencies
        freqs = [
            abs(round(f))
            for f in obj.properties.vibrations.frequencies[:num_imag_modes]
        ]

        freqs_str = ", ".join(f"{f}<i>i</i>" for f in freqs)
        write_str += f"<b><i>ν<sub>imag</sub></i></b> = {num_imag_modes} ({freqs_str}) cm<sup>–1</sup>"

    return write_str


class SIWriters(Enum):
    """
    Enum for the different types of SI writers.
    """

    HYDROGEN_BOND = _hydrogen_bond_format


class SI:
    def __init__(
        self,
        path: str | pl.Path,
        append_mode: bool = False,
        font: str = "Arial",
        format: SIWriters = SIWriters.HYDROGEN_BOND,
    ):
        """
        Class for creating supporting information (SI) files in Microsoft Word.

        Args:
                path: the location of the Word file. Does not have to have a file-extension.
                append_mode: whether to append to or overwrite the file.
        """

        self.path = pl.Path(path).with_suffix(".docx")
        self.doc = docx.Document()
        self._format_writer = format

        if append_mode:
            try:
                self.doc = docx.Document(self.path)
            except FileNotFoundError:
                pass

        # set the font to Calibri
        self.doc.styles["Normal"].font.name = font

    def __enter__(self):
        return self

    def __exit__(self, *args, **kwargs):
        self.doc.save(self.path)

    def add_xyz(self, obj: str | Result, title: str):
        """
        Add the coordinates and information about a calculation to the SI.
        It will add the electronic bond energy, Gibb's free energy, enthalpy and imaginary mode, as well as the coordinates of the molecule.

        Args:
                obj: a string specifying a calculation directory or a `TCutility.results.Result` object from a calculation.
                title: title to be written before the coordinates and information.
        """
        # if isinstance(obj, str):
        #     obj = results.read(obj)

        # # title is always bold
        # s = f"<b>{title}</b><br>"

        # parser = HtmlToDocx()

        # # add the electronic energy, Gibbs free energy, enthalpy and imaginary mode
        # s = self._format_writer(obj, s)

        # # remove trailing line breaks
        # s = s.removesuffix("<br>")

        # # coords should be written in mono-type font with 8 decimals and 4 spaces between each coordinate
        # s += "<pre>"
        # for atom in obj.molecule.output:
        #     s += f"{atom.symbol:2}    {atom.coords[0]: .8f}    {atom.coords[1]: .8f}    {atom.coords[2]: .8f}<br>"
        # s += "</pre>"
        # parser.add_html_to_document(s, self.doc)
        parser = HtmlToDocx()
        parser.add_html_to_document(xyz.Default()(obj), self.doc)
        # return xyz.Default()(obj)

    def add_heading(self, text: str, level: int = 1):
        """
        Add a heading to the file. This method has the same arguments and functionality as docx.Document.add_heading.
        """
        self.doc.add_heading(text, level)
