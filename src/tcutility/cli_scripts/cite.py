"""Module containing functions for generating citations"""

from importlib.util import find_spec

from tcutility import errors

if find_spec("docx") is None:
    raise errors.MissingOptionalPackageError("docx")
if find_spec("htmldocx") is None:
    raise errors.MissingOptionalPackageError("htmldocx")


import os
from math import ceil
from typing import List

import click
import docx
import htmldocx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from tcutility import cite, spell_check
from tcutility.data import functionals


class Docx:
    """
    Small class that handles writing to a docx file. This should and will be moved to its own module in TCutility soon.
    """

    def __init__(self, file="test.docx", overwrite=False):
        self.file = file
        if not os.path.exists(file) or overwrite:
            self.doc = docx.Document()
        else:
            self.doc = docx.Document(file)

        self.doc.styles["Normal"].font.name = "Times New Roman"
        self.doc.styles["Normal"].font.size = Pt(12)
        self.doc.styles["Normal"].paragraph_format.space_after = 0
        self.html_parser = htmldocx.HtmlToDocx()

    def __enter__(self):
        return self

    def __exit__(self, *args):
        self.doc.save(self.file)

    def write_paragraph(self, text, alignment=WD_ALIGN_PARAGRAPH.LEFT):
        """
        Write a piece of text as a pragraph to this Docx file.
        This function will parse any HTML that is given in the text.
        E.g. you can use the <b></b> tags to make a piece of text bold.
        """
        self.html_parser.add_html_to_document(text, self.doc)
        self.doc.paragraphs[-1].alignment = alignment

    def open(self):
        """
        Open this file in Word.
        """
        os.system(f"open {self.file}")


program_references = {
    "ams": [
        "10.1002/jcc.1056",
        "10.1007/s002140050353",
    ],
    "adf": [
        "10.1002/jcc.1056",
        "10.1007/s002140050353",
    ],
    "orca": [
        "10.1002/wcms.81",
        "10.1063/5.0004608",
        "10.1002/wcms.1606",
    ],
    "dftb": [],
    "xtb": [],
    "cosmo": [],
    "crest": [],
    "pyfrag": [
        "10.1002/jcc.20786",
        "10.1002/jcc.25871",
        "10.5281/zenodo.1045523",
    ],
    "pyorb": [],
    "cylview": [],
}

methodology_references = {
    "fmatsfo": [],
    "eda": ["10.1002/9780470125922.ch1", "10.1515/9783110660074-008"],
    "asm": [
        "10.1038/s41596-019-0265-0",
        "10.1002/(SICI)1096-987X(19990115)20:1<114::AID-JCC12>3.0.CO;2-L",
        "10.1039/B926828F",
        "10.1039/C4CS00055B",
        "10.1002/wcms.1221",
        "10.1002/anie.201701486",
        "10.1039/D1CC02042K",
    ],
    "zora": [
        "10.1063/1.466059",
        "10.1063/1.467943",
    ],
    "ksmo": [],
    "vdd": [
        "10.1002/jcc.10351",
        "10.1002/jcc.27184",
        "10.1039/c5cp07483e",
    ],
    "hydrogen-bonding": [],
    "halogen-bonding": [],
    "chalcogen-bonding": [],
    "pnictogen-bondding": [],
    "zlm-fit": [
        "10.1021/ct500172n",
    ],
    "becke-grid": [
        "1.1002/jcc.23323",
    ],
    "sto": [
        "10.1002/jcc.10255",
    ],
    "vibrational-analysis": [
        "10.1016/s0010-4655(96)00120-8",
        "10.1016/S0010-4655(96)00119-1",
        "10.1002/qua.20653",
    ],
    "irc": [],
    "d3": ["10.1063/1.3382344"],
    "d3bj": ["10.1002/jcc.21759"],
    "d4": ["10.1063/1.5090222"],
}


doi_order = []


def format_paragraph(dois, style):
    paragraphs = []
    for doi in dois:
        if doi not in doi_order:
            doi_order.append(doi)
        try:
            citation = cite.cite(doi, style=style, mode="plain")
        except Exception as exp:
            print("\t" + str(exp))
            raise exp

        print(f"  [{doi}] {citation}")
        paragraph = f"[{doi_order.index(doi) + 1}] {cite.cite(doi, style=style, mode='html')}"
        paragraphs.append(paragraph)

    return paragraphs


def _print_rect_list(printables, spaces_before=0):
    """
    This function prints a list of strings in a rectangle to the output.
    This is similar to what the ls program does in unix.
    """
    n_shell_col = os.get_terminal_size().columns
    prev_mat = []
    prev_col_lens = []

    # we first have to determine the correct dimensions of our rectangle
    for ncol in range(1, n_shell_col):
        # the number of rows for the number of columns
        nrows = ceil(len(printables) / ncol)
        # we then get what the rectangle would be
        mat = [printables[i * ncol : (i + 1) * ncol] for i in range(nrows)]
        # and determine for each column the width
        col_lens = [max([len(row[i]) for row in mat if i < len(row)] + [0]) for i in range(ncol)]
        # then calculate the length of each row based on the column lengths
        # we use a spacing of 2 spaces between each column
        row_len = spaces_before + sum(col_lens) + 2 * len(col_lens) - 2

        # if the rows are too big we exit the loop
        if row_len > n_shell_col:
            break

        # store the previous loops results
        prev_col_lens = col_lens
        prev_mat = mat

    # then print the strings with the right column widths
    for row in prev_mat:
        print(" " * spaces_before + "  ".join([x.ljust(col_len) for x, col_len in zip(row, prev_col_lens)]))


@click.command("cite")
@click.argument("objects", type=str, nargs=-1, required=False)
@click.option("-w", "--wiley", help="Set the citation style to Wiley. This is the default style.", flag_value="wiley", default=True)
@click.option("-a", "--acs", help="Set the citation style to ACS.", flag_value="acs")
@click.option("-r", "--rsc", help="Set the citation style to RSC.", flag_value="rsc")
@click.option("-o", "--output", help="The output Word file to write the citations to.", type=str, default="citations.docx")
@click.option("-l", "--list_citations", help="List currently available citations.", is_flag=True, default=False)
def generate_citations(objects: List[str], wiley: bool, acs: bool, rsc: bool, output: str, list_citations: bool) -> None:
    """
    Generate citations for objects.

    Currently supports generating citations for functionals, basis-sets, programs, methodologies and DOIs.
    This program also generates, and if possible, opens a Word document that contains the formatted citations.
    Multiple objects can be given separated by spaces.
    If the supplied object is also a file path it will read each line as a separate object.

    Example usage:

    > tc cite ADF
    Program ORCA
      [10.1002/wcms.81] F. Neese, WIREs Comput. Mol. Sci. 2011, 2, 73-78.
      [10.1063/5.0004608] F. Neese, F. Wennmohs, U. Becker, C. Riplinger, J. Chem. Phys. 2020, 152.
      [10.1002/wcms.1606] F. Neese, WIREs Comput. Mol. Sci. 2022, 12.

    > tc cite BP86 BLYP OLYP OPBE D3BJ
    Functional BP86
      [10.1103/PhysRevA.38.3098] A. D. Becke, Phys. Rev. A 1988, 38, 3098-3100.
      [10.1103/PhysRevB.33.8800] J. P. Perdew, W. Yue, Phys. Rev. B 1986, 33, 8800-8802.
    Functional BLYP
      [10.1103/PhysRevA.38.3098] A. D. Becke, Phys. Rev. A 1988, 38, 3098-3100.
      [10.1103/PhysRevB.37.785] C. Lee, W. Yang, R. G. Parr, Phys. Rev. B 1988, 37, 785-789.
    Functional OLYP
      [10.1080/00268970010018431] N. C. Handy, A. J. Cohen, Mol. Phys. 2001, 99, 403-412.
      [10.1103/PhysRevB.37.785] C. Lee, W. Yang, R. G. Parr, Phys. Rev. B 1988, 37, 785-789.
    Functional OPBE
      [10.1080/00268970010018431] N. C. Handy, A. J. Cohen, Mol. Phys. 2001, 99, 403-412.
      [10.1103/PhysRevLett.77.3865] J. P. Perdew, K. Burke, M. Ernzerhof, Phys. Rev. Lett. 1996, 77, 3865-3868.
    Methodology D3BJ
      [10.1002/jcc.21759] S. Grimme, S. Ehrlich, L. Goerigk, J. Comput. Chem. 2011, 32, 1456-1465.
    """
    available_citations = list(program_references.keys()) + list(methodology_references.keys()) + list(functionals.functionals.keys())
    if list_citations:
        print("OBJECTS WITH AVAILABLE REFERENCES:")
        print("    Programs")
        print("    ========")
        printables = [prog for prog, dois in program_references.items() if len(dois) > 0]
        _print_rect_list(printables, 4)
        print()

        print("    Methodology")
        print("    ===========")
        printables = [meth for meth, dois in methodology_references.items() if len(dois) > 0]
        _print_rect_list(printables, 4)
        print()

        print("    Functionals")
        print("    ===========")
        printables = [xc_info.path_safe_name for xc, xc_info in functionals.functionals.items() if len(xc_info.dois) > 0]  # type: ignore
        _print_rect_list(printables, 4)
        print()

        return

    if len(objects) == 1 and os.path.isfile(objects[0]):
        with open(objects[0]) as inp:
            objects = [line.strip() for line in inp.readlines()]

    style = "wiley" if wiley else "acs" if acs else "rsc"

    with Docx(file=output, overwrite=True) as out:
        for obj in objects:
            paragraphs = None  # try to format a functional
            paragraph_title = ""
            try:
                xc_info = functionals.get_functional(obj)
                print("Functional", obj)
                paragraph_title = f"XC-Functional: <b>{xc_info.name_html}</b>"
                paragraphs = format_paragraph(xc_info.dois, style=style)

            except KeyError:
                pass

            # if its not a functional we look in the programs
            if obj.lower() in program_references:
                print("Program", obj)
                paragraph_title = f"Program: <b>{obj}</b>"
                paragraphs = format_paragraph(program_references[obj.lower()], style=style)

            # and the methodologies
            if obj.lower() in methodology_references:
                print("Methodology", obj)
                paragraph_title = f"Method: <b>{obj}</b>"
                paragraphs = format_paragraph(methodology_references[obj.lower()], style=style)

            # if we still dont have a paragraphs we check if it is a DOI
            if paragraphs is None and obj.startswith("10."):
                print("DOI")
                paragraph_title = f"DOI: <b>{obj}</b>"
                paragraphs = format_paragraph([obj], style=style)

            if paragraphs is None:
                spell_check.make_suggestion(obj, available_citations, ignore_case=True)
                continue

            out.write_paragraph(paragraph_title)
            for paragraph in paragraphs:
                out.write_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            out.write_paragraph(" ")

    out.open()
