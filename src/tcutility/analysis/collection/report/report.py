import pathlib as pl

import docx
from htmldocx import HtmlToDocx
from tcutility.analysis.collection.report.formatters.generic import WordFormatter
from tcutility.analysis.collection.report.formatters.xyz import XYZFormatter
from tcutility.results import Result


class SI:
    def __init__(self, path: str | pl.Path, append_mode: bool = False, font: str = "Arial", format: WordFormatter = XYZFormatter):
        """Initializes the SI class for creating supporting information (SI) files in Microsoft Word format.

        This class is responsible for creating and managing a Microsoft Word document that serves as supporting information (SI) for reports or publications. It allows for the addition of various elements such as text, headings, and formatted content from HTML.

        Args:
            path (str | pl.Path): The location of the Word file. Does not have to have a file-extension.
            append_mode (bool, optional): Whether to append to or overwrite the file. Defaults to False.
            font (str, optional): The font to be used in the document. Defaults to "Arial".
            format (WordFormatter, optional): The formatter to be used for adding content to the document. Defaults to XYZFormatter.

        Attributes:
            path (pl.Path): The path to the Word document.
            doc (docx.Document): The Word document object.
            _format_writer (WordFormatter): The formatter instance for adding formatted content.
        """
        self.path = pl.Path(path).with_suffix(".docx")
        self.doc = docx.Document()
        self._format_writer = format

        if append_mode:
            try:
                self.doc = docx.Document(str(self.path))
            except FileNotFoundError:
                pass

        # Set the font to the specified font
        self.doc.styles["Normal"].font.name = font

    def __enter__(self):
        """Enables the use of the class as a context manager."""
        return self

    def __exit__(self, *args, **kwargs):
        """Saves the document upon exiting the context manager."""
        self.doc.save(self.path)

    def add_xyz(self, obj: str | Result, title: str | None) -> str:
        """Adds XYZ formatted content to the document.

        This method is responsible for adding the coordinates and information about a calculation to the supporting information document. It includes details such as the electronic bond energy, Gibb's free energy, enthalpy, imaginary mode, and the coordinates of the molecule.

        Args:
            obj (str | Result): A string specifying a calculation directory or a `TCutility.results.Result` object from a calculation.
            title (str | None): The title to be written before the coordinates and information. If None, no title is added.

        Returns:
            str: The formatted content that was added to the document.
        """
        # Add title to the document
        if title is not None:
            self.add_heading(title)

        parser = HtmlToDocx()
        parser.add_html_to_document(self._format_writer.write(obj), self.doc)
        return self._format_writer.write(obj)

    def add_heading(self, text: str, level: int = 1):
        """Adds a heading to the document.

        This method allows for the addition of a heading to the Word document, with customizable text and level.

        Args:
            text (str): The text of the heading.
            level (int, optional): The level of the heading (determines the size and style). Defaults to 1.
        """
        self.doc.add_heading(text, level)


def main(): ...


if __name__ == "__main__":
    main()
