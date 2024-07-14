import os
import pathlib as pl
from typing import List, Union

import docx
from htmldocx import HtmlToDocx
from tcutility.analysis.report.formatters.generic import WordFormatter
from tcutility.analysis.report.formatters.xyz import XYZFormatter
from tcutility.results import Result, read


class SI:
    def __init__(self, path: Union[str, pl.Path], append_mode: bool = False, font: str = "Arial", format: WordFormatter = XYZFormatter()) -> None:
        """Initializes the SI class for creating supporting information (SI) files in Microsoft Word format.

        This class is responsible for creating and managing a Microsoft Word document that serves as supporting information (SI) for reports or publications.
        It allows for the addition of various elements such as text, headings, and formatted content from HTML.

        Args:
            path (str | pl.Path): The location of the Word file. Does not have to have a file-extension.
            append_mode (bool, optional): Whether to append to or overwrite the file. Defaults to False.
            font (str, optional): The font to be used in the document. Defaults to "Arial".
            format (WordFormatter, optional): The formatter to be used for adding content to the document. Defaults to XYZFormatter.

        Attributes:
            path (pl.Path): The path to the Word document.
            doc (docx.Document): The Word document object.
            _format_writer (WordFormatter): The formatter instance for adding formatted content.
        """
        self.path = pl.Path(path).with_suffix(".docx")
        self.doc = docx.Document()
        self._format_writer = format

        if append_mode:
            try:
                self.doc = docx.Document(str(self.path))
            except FileNotFoundError:
                pass
            except docx.opc.exceptions.PackageNotFoundError:  # type: ignore  # opc not found in docx by mypy
                pass

        # Set the font to the specified font
        self.doc.styles["Normal"].font.name = font

    def __enter__(self):
        """Enables the use of the class as a context manager."""
        return self

    def __exit__(self, *args, **kwargs):
        """Saves the document upon exiting the context manager."""
        self.doc.save(self.path)

    def add_xyz(self, obj: Union[str, Result], title: Union[str, None] = None) -> None:
        """Adds XYZ formatted content to the document.

        This method is responsible for adding the coordinates and information about a calculation to the supporting information document.
        It includes details such as the electronic bond energy, Gibb's free energy, enthalpy, imaginary mode, and the coordinates of the molecule.

        Args:
            obj (str | Result): A string specifying a calculation directory or a `TCutility.results.Result` object from a calculation.
            title (str | None): The title to be written before the coordinates and information. If None, no title is added.

        Returns:
            None
        """
        ret_str = ""

        # Add the formatted content to the document
        if isinstance(obj, str):
            obj = read(obj)

        ret_str += self._format_writer.format(obj, title=title)

        # print(ret_str)
        parser = HtmlToDocx()
        parser.add_html_to_document(ret_str, self.doc)
        return

    def add_heading(self, text: str, level: int = 1) -> None:
        """Adds a heading to the document.

        This method allows for the addition of a heading to the Word document, with customizable text and level.

        Args:
            text (str): The text of the heading.
            level (int, optional): The level of the heading (determines the size and style). Defaults to 1.
        """
        self.doc.add_heading(text, level)

def get_subdirs(root_folder: pl.Path) -> List[pl.Path]:
    """Iteratively searches through a folder and returns all the most nested subdirs."""
    most_nested_subdirs = []
    for root, dirs, files in os.walk(root_folder):
        # If 'dirs' is empty, it means 'root' contains no subdirectories, thus it is most nested.
        if not dirs:
            most_nested_subdirs.append(pl.Path(root))
    return most_nested_subdirs

def replace_files_rkf_to_ams_rkf(root_folder: pl.Path) -> None:
    """ Iteratively searches through a folder and replaces all files with the extension '.rkf' to '.ams.rkf', except if the file has 'adf.rkf' in the name."""
    for root, dirs, files in os.walk(root_folder):
        for file in files:
            if file.endswith(".rkf") and "adf.rkf" not in file:
                new_name = file.replace(".rkf", ".ams.rkf")
                os.rename(pl.Path(root) / file, pl.Path(root) / new_name)

def main():
    from tcutility.results import read

    calc_dir = pl.Path("__file__").resolve().parents[0] / "test" / "fixtures"
    main_path = pl.Path("__file__").resolve().parents[0] / "examples"

    # all_subdirs = [folder for folder in calc_dir.iterdir() if folder.is_dir()]
    all_subdirs = get_subdirs(calc_dir)
    res_objects = []
    for dir_ in all_subdirs:
        try:  # Try to read the results of the calculation
            res_objects.append(read(dir_))
        except Exception:
            pass

    with SI(main_path / "test", append_mode=False) as si:
        si.add_heading("SI Electronegativity project")
        for obj in res_objects:
            si.add_xyz(obj=obj)


if __name__ == "__main__":
    main()
