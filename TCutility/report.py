from TCutility import results, formula
import docx
from htmldocx import HtmlToDocx


class SI:

	def __init__(self, path: str, append_mode: bool = False):
		'''
		Class for creating supporting information (SI) files in Microsoft Word.

		Args:
			path: the location of the Word file. Does not have to have a file-extension.
			append_mode: whether to append to or overwrite the file.
		'''

		self.path = path.removesuffix('.docx') + '.docx'
		self.doc = docx.Document()

		self.doc.styles['Normal'].font.name = 'Calibri'
		self.append_mode = append_mode

	def __enter__(self):
		return self

	def __exit__(self, *args, **kwargs):
		self.doc.save(self.path)

	def add_xyz(self, obj: str or dict, title: str):
		'''
		Add the coordinates and information about a calculation to the SI. 
		It will add the electronic bond energy, Gibb's free energy, enthalpy and imaginary mode, as well as the coordinates of the molecule.

		Args:
			obj: a string specifying a calculation directory or a `TCutility.results.Result` object from a calculation.
			title: title to be written before the coordinates and information.
		'''
		if isinstance(obj, str):
			obj = results.read(obj)

		# title is always bold
		s = f'<b>{title}</b><br>'

		parser = HtmlToDocx()

		# add electronic energy. E should be bold and italics. Unit will be kcal mol^-1
		E = str(round(obj.properties.energy.bond, 1)).replace('-', '—')
		s += f'<b><i>E</i></b> = {E} kcal mol<sup>—1</sup><br>'

		# add Gibbs and enthalpy if we have them
		if obj.properties.energy.gibbs:
			G = str(round(obj.properties.energy.gibbs, 1)).replace('-', '—')
			s += f'<b><i>G</i></b> = {G} kcal mol<sup>—1</sup><br>'
		if obj.properties.energy.enthalpy:
			H = str(round(obj.properties.energy.enthalpy, 1)).replace('-', '—')
			s += f'<b><i>H</i></b> = {H} kcal mol<sup>—1</sup><br>'

		# add imaginary frequency if we have one
		if obj.properties.vibrations:
			if obj.properties.vibrations.number_of_imag_modes == 1:
				freq = abs(round(obj.properties.vibrations.frequencies[0]))
				s += f'<b><i>ν<sub>imag</sub></i></b> = {freq}<i>i</i> cm<sup>—1</sup>'
		
		# remove trailing line breaks
		s = s.removesuffix('<br>')
		
		# coords should be written in mono-type font with 8 decimals and 4 spaces between each coordinate
		s += '<pre>'
		for atom in obj.molecule.output:
			s += f'{atom.symbol:2}    {atom.coords[0]: .8f}    {atom.coords[1]: .8f}    {atom.coords[2]: .8f}<br>'
		s += '</pre>'
		parser.add_html_to_document(s, self.doc)

	def add_heading(self, text: str, level: int = 1):
		'''
		Add a heading to the file. This method has the same arguments and functionality as docx.Document.add_heading.
		'''
		self.doc.add_heading(text, level)


if __name__ == '__main__':
	with SI('test.docx') as si:
		si.add_heading('Test molecules:')
		si.add_xyz('/Users/yumanhordijk/PhD/TheoCheM_stack/TCutility/test/fixtures/level_of_theory/M06_2X', formula.molecule('C2H6'))
		si.add_xyz('/Users/yumanhordijk/PhD/TheoCheM_stack/TCutility/test/fixtures/ethanol', formula.molecule('C2H5OH'))
